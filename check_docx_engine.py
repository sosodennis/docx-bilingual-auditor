import sys
import os
import re
import html  # Added for HTML escaping
import json  # Added for safe JS string generation
import pandas as pd
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from thefuzz import fuzz  # [RESTORED] Re-imported for fuzzy matching

if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

chi_input_docx = os.path.join(BASE_DIR, "chi_input.docx")
eng_input_docx = os.path.join(BASE_DIR, "eng_input.docx")
output_html = os.path.join(BASE_DIR, "report.html")

# Set pandas display options
pd.set_option('display.unicode.east_asian_width', True)
pd.set_option('display.max_colwidth', None) # Changed to None to allow full HTML rendering

# ==========================================
# 1. Data Structures & Utility Functions
# ==========================================

class DocumentSection:
    def __init__(self, title):
        self.title = title
        self.content_blocks = []
        # Stores list of dicts: [{'text': 'BoldWord', 'context': 'Full sentence containing BoldWord'}]
        self.bold_data = []      
        self.underline_data = [] 
    
    def add_content(self, text):
        if text.strip():
            self.content_blocks.append(text)
            
    def add_bold_items(self, item_list):
        """item_list expects [{'text':..., 'context':...}, ...]"""
        if item_list:
            self.bold_data.extend(item_list)

    def add_underline_items(self, item_list):
        """item_list expects [{'text':..., 'context':...}, ...]"""
        if item_list:
            self.underline_data.extend(item_list)

    def get_full_content(self):
        return "\n".join(self.content_blocks)

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_table_text(table):
    rows_text = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        rows_text.append(" | ".join(row_data))
    return "\n[表格開始]\n" + "\n".join(rows_text) + "\n[表格結束]"

def extract_bold_items(block):
    """
    Extract bold text AND its context (the full paragraph text).
    Returns a list of dicts: [{'text': '...', 'context': '...'}]
    """
    items = []
    
    def process_paragraph(para):
        para_text = para.text.strip() # The context is the full paragraph
        if not para_text:
            return

        buffer_text = ""
        for run in para.runs:
            # Check for bold explicitly
            if run.bold:
                buffer_text += run.text
            else:
                if buffer_text.strip():
                    items.append({
                        'text': buffer_text.strip(), 
                        'context': para_text
                    })
                buffer_text = ""
        # Flush buffer at end of paragraph
        if buffer_text.strip():
            items.append({
                'text': buffer_text.strip(), 
                'context': para_text
            })

    if isinstance(block, Paragraph):
        process_paragraph(block)
    elif isinstance(block, Table):
        for row in block.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)
                    
    return items

def extract_underline_items(block):
    """
    Extract underlined text AND its context.
    Returns a list of dicts.
    """
    items = []
    
    def process_paragraph(para):
        para_text = para.text.strip()
        if not para_text:
            return

        buffer_text = ""
        for run in para.runs:
            if run.underline:
                buffer_text += run.text
            else:
                if buffer_text.strip():
                    items.append({
                        'text': buffer_text.strip(), 
                        'context': para_text
                    })
                buffer_text = ""
        if buffer_text.strip():
            items.append({
                'text': buffer_text.strip(), 
                'context': para_text
            })

    if isinstance(block, Paragraph):
        process_paragraph(block)
    elif isinstance(block, Table):
        for row in block.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)
                    
    return items

# ==========================================
# 2. Core Processing Logic
# ==========================================

def parse_document_sections(file_path, toc_keyword, regex_pattern, log_func=print):
    if not os.path.exists(file_path):
        log_func(f"Error: File not found -> {file_path}")
        return []

    try:
        doc = Document(file_path)
    except Exception as e:
        log_func(f"Error reading docx: {e}")
        return []
    
    # --- DEBUG START ---
    log_func(f"\n[DEBUG] Start analyzing file: {os.path.basename(file_path)}")
    # --- DEBUG END ---
    
    # Step A: Extract TOC
    toc_titles = []
    for i, table in enumerate(doc.tables):
        try:
            sample_rows = table.rows[:5]
            header_check = "".join([c.text.strip() for r in sample_rows for c in r.cells])
        except Exception:
            header_check = ""

        if toc_keyword in header_check:
            for row in table.rows:
                cells = row.cells
                if not cells: continue
                
                # [FIX v2]: Multi-line Title Extraction
                full_cell_text = cells[0].text.strip()
                lines = [line.strip() for line in full_cell_text.split('\n') if line.strip()]
                
                title_parts = []
                capture_mode = False
                
                for line in lines:
                    if re.search(regex_pattern, line, re.IGNORECASE):
                        capture_mode = True
                        title_parts.append(line)
                    elif capture_mode:
                        if len(line) > 80: 
                            break
                        title_parts.append(line)
                    else:
                        continue
                
                if title_parts:
                    full_title = " ".join(title_parts)
                    toc_titles.append(full_title)
            break

    extracted_sections = []
    current_section = None
    
    if not toc_titles:
        log_func(f"Warning: No TOC found in {os.path.basename(file_path)}.")
        log_func(">>> Switching to 'Whole Document' mode (Single Section).")
        single_section = DocumentSection(f"Whole Document ({os.path.basename(file_path)})")
        extracted_sections = [single_section]
        current_section = single_section
    else:
        log_func(f"Found {len(toc_titles)} sections in {os.path.basename(file_path)}")
        extracted_sections = []
        current_section = None

    # Step B: Full text scan
    title_cursor = 0 
    
    for block in iter_block_items(doc):
        block_text_for_check = ""
        full_block_content = ""
        
        current_bold_items = extract_bold_items(block)
        current_underline_items = extract_underline_items(block)

        if isinstance(block, Paragraph):
            block_text_for_check = block.text.strip()
            full_block_content = block_text_for_check
            
        elif isinstance(block, Table):
            if block.rows:
                block_text_for_check = " ".join([c.text.strip() for c in block.rows[0].cells])
            full_block_content = get_table_text(block)

        if not block_text_for_check:
            continue

        matched_new_title = False
        if toc_titles and title_cursor < len(toc_titles):
            target_title = toc_titles[title_cursor]
            
            # [HYBRID STRATEGY]
            # 1. Prepare strings
            target_clean = "".join(target_title.split())
            block_clean = "".join(block_text_for_check.split())
            
            # 2. Attempt Exact Match (Best case)
            if target_clean and block_clean and target_clean.lower() == block_clean.lower():
                matched_new_title = True
            
            # 3. Attempt Fuzzy Match (Robustness for Certificate vs Confirmation)
            # Only if Exact Match failed
            elif target_title and block_text_for_check:
                len_t = len(target_clean)
                len_b = len(block_clean)
                
                # Length Safety Check:
                # The block text length must be within +/- 20% of the target title length.
                # This prevents matching a title to a long paragraph that happens to contain similar words.
                if len_t > 0:
                    ratio = len_b / len_t
                    if 0.8 <= ratio <= 1.2:
                        # Use fuzz.ratio for strict similarity check
                        # Threshold 80 should cover "Certificate" vs "Confirmation" (~85% similarity in context)
                        score = fuzz.ratio(target_title.lower(), block_text_for_check.lower())
                        if score >= 80:
                            matched_new_title = True
        
        if matched_new_title:
            if current_section:
                extracted_sections.append(current_section)
            clean_title = toc_titles[title_cursor]
            current_section = DocumentSection(clean_title)
            title_cursor += 1
            # Don't 'continue' here, allow the title itself to be processed if needed

        if current_section:
            current_section.add_content(full_block_content)
            current_section.add_bold_items(current_bold_items)
            current_section.add_underline_items(current_underline_items)

    if toc_titles and current_section:
        extracted_sections.append(current_section)

    return extracted_sections

# ==========================================
# 3. HTML Generation (Enhanced with Clipboard)
# ==========================================

def generate_html_report(sections_chi, sections_eng, output_path):
    # CSS & JS for Clipboard Functionality
    HTML_HEADER = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
    <meta charset="UTF-8">
    <title>Comparison Report</title>
    <style>
        body{font-family:"Microsoft JhengHei",Arial,sans-serif;background-color:#f4f4f9;margin:40px}
        h1{text-align:center;color:#333}
        .section-container{background:#fff;padding:25px;margin-bottom:30px;border-radius:8px;box-shadow:0 4px 6px rgba(0,0,0,.1)}
        .section-header{border-bottom:2px solid #3498db;padding-bottom:15px;margin-bottom:20px;display:flex;justify-content:space-between;align-items:center}
        .section-title{font-size:1.4em;font-weight:700;color:#2c3e50}
        .sub-info{color:#7f8c8d;font-size:.95em;margin-top:5px}
        .category-header{margin-top:30px;margin-bottom:10px;font-size:1.1em;font-weight:700;color:#2980b9;border-left:5px solid #2980b9;padding-left:10px;background-color:#ecf0f1;padding:8px}
        
        table{width:100%;border-collapse:collapse;table-layout:fixed;margin-bottom:20px}
        th,td{border:1px solid #e0e0e0;padding:10px;text-align:left;vertical-align:top;word-wrap:break-word}
        thead th{background-color:#3498db;color:#fff;font-weight:700}
        
        /* New Styles for Copy Button */
        .item-wrapper { display: flex; align-items: flex-start; gap: 8px; }
        .copy-btn {
            background-color: #eee; border: 1px solid #ccc; border-radius: 4px;
            cursor: pointer; padding: 2px 6px; font-size: 14px; flex-shrink: 0;
            transition: background-color 0.2s;
        }
        .copy-btn:hover { background-color: #ddd; }
        .copy-btn:active { background-color: #bbb; }
        .text-content { font-weight: 500; }
        
        tr>*:nth-child(1){width:50px;text-align:center;background-color:#f8f9fa;color:#666;font-weight:700}
        tr:nth-child(even){background-color:#fcfcfc}
        tr:hover{background-color:#f1f8ff;transition:.2s}
        .empty-msg{color:#95a5a6;font-style:italic;padding:10px;border:1px dashed #ccc;background:#fafafa}
    </style>
    <script>
        function copyToClipboard(text, btnElement) {
            navigator.clipboard.writeText(text).then(function() {
                // Logic to handle double clicks: Clear existing timer if present
                if (btnElement.dataset.timer) {
                    clearTimeout(parseInt(btnElement.dataset.timer));
                }

                // Visual feedback: Flash color only, do not change text
                btnElement.style.backgroundColor = "#90EE90"; // Light green

                // Revert color after 500ms
                const timerId = setTimeout(() => {
                    btnElement.style.backgroundColor = "#eee"; // Revert to original
                    delete btnElement.dataset.timer; // Clean up
                }, 500);

                // Store timer ID on the element so we can cancel it if clicked again
                btnElement.dataset.timer = timerId;

            }, function(err) {
                console.error('Could not copy text: ', err);
                alert("Failed to copy context.");
            });
        }
    </script>
    </head>
    <body>
    <h1>Chinese-English Document Comparison Report</h1>
    """
    
    HTML_FOOTER = """
    <div style="text-align:center;color:#aaa;margin-top:50px;font-size:.8em;border-top:1px solid #eee;padding-top:20px">
        Generated by Docx Auditor GUI
    </div>
    </body></html>
    """

    def format_item_html(item_dict):
        """
        Takes a dict {'text': '...', 'context': '...'} and returns HTML string with button.
        """
        if not item_dict:
            return ""
        
        # Escape HTML special characters for display
        # Also clean up newlines in the *displayed* text to avoid weird spacing
        text_val = item_dict.get('text', '')
        display_text = html.escape(text_val).replace('\n', ' ')
        
        # Escape context for insertion into JavaScript string
        # We use json.dumps to safely encode the string for JS, then slice off the quotes
        raw_context = item_dict.get('context', '')
        js_safe_context = json.dumps(raw_context).replace('"', '&quot;')
        
        # HTML Block - Constructed in one line to avoid introducing \n into the output
        html_block = (
            f'<div class="item-wrapper">'
            f'<button class="copy-btn" onclick="copyToClipboard({js_safe_context}, this)" title="Copy context to search">📋</button>'
            f'<span class="text-content">{display_text}</span>'
            f'</div>'
        )
        return html_block

    def create_comparison_table(list_c, list_e, col_c_name, col_e_name):
        # list_c and list_e are lists of dictionaries now
        if not list_c and not list_e:
            return '<div class="empty-msg">No Content</div>'
            
        max_rows = max(len(list_c), len(list_e))
        
        # Convert dictionaries to HTML strings
        html_c = [format_item_html(item) for item in list_c]
        html_e = [format_item_html(item) for item in list_e]
        
        # Pad lists
        html_c_padded = html_c + [""] * (max_rows - len(html_c))
        html_e_padded = html_e + [""] * (max_rows - len(html_e))
        
        df = pd.DataFrame({
            col_c_name: html_c_padded, 
            col_e_name: html_e_padded
        })
        df.index += 1
        
        # escape=False is CRITICAL here to render the button HTML we created
        return df.to_html(classes='table', border=0, justify='left', escape=False)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(HTML_HEADER)
        max_sections = max(len(sections_chi), len(sections_eng))
        
        for i in range(max_sections):
            sec_c = sections_chi[i] if i < len(sections_chi) else None
            sec_e = sections_eng[i] if i < len(sections_eng) else None
            
            title_c = sec_c.title if sec_c else "(No such section)"
            title_e = sec_e.title if sec_e else "(Section Missing)"
            
            f.write(f'<div class="section-container">')
            f.write(f'<div class="section-header"><div><strong>Section {i+1}</strong></div></div>')
            f.write(f'<div class="sub-info"><b>CH Title:</b> {title_c}</div>')
            f.write(f'<div class="sub-info"><b>EN Title:</b> {title_e}</div>')
            
            f.write('<div class="category-header">1. Bold Text (Click 📋 to copy context)</div>')
            f.write(create_comparison_table(
                sec_c.bold_data if sec_c else [], 
                sec_e.bold_data if sec_e else [], 
                "Chinese (Bold)", 
                "English (Bold)"
            ))
            
            f.write('<div class="category-header">2. Underlined Text (Click 📋 to copy context)</div>')
            f.write(create_comparison_table(
                sec_c.underline_data if sec_c else [], 
                sec_e.underline_data if sec_e else [], 
                "Chinese (Underline)", 
                "English (Underline)"
            ))
            
            f.write('</div>') # End section-container
            
        f.write(HTML_FOOTER)